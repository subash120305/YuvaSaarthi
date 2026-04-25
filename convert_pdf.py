from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

def main():
    pdf_path = "YSR.pdf"
    print(f"Reading {pdf_path}...")
    try:
        reader = PdfReader(pdf_path)
        text_content = ""
        for i, page in enumerate(reader.pages):
            text_content += page.extract_text() + "\n"
        print(f"Extracted {len(text_content)} characters.")
    except Exception as e:
        print(f"Failed to read PDF: {e}")
        text_content = "Failed to extract text from YSR.pdf. Please append manually."

    doc = Document()
    
    # Setup styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    # IEEE Title
    title = doc.add_heading("YuvaSaarthi: AI powered educational assistance chatbot", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Authors
    doc.add_paragraph()  # spacing
    auth_para = doc.add_paragraph()
    auth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    authors = [
        "Subash S (subashsxyz@gmail.com)",
        "S A Shanjith (sashanjith@gmail.com)",
        "Yashas P (yashaspk98@gmail.com)",
        "Harshith K N (harshithkn04@gmail.com)",
        "Prof. Mahima M Gowda (mahima.gowda.m@reva.edu.in)"
    ]
    
    for author in authors:
        run = auth_para.add_run(author + "\n")
        run.font.size = Pt(11)
        
    uni_run = auth_para.add_run("REVA University, Bangalore")
    uni_run.italic = True
    uni_run.font.size = Pt(10)
    
    doc.add_paragraph() # spacing
    
    # Body
    doc.add_heading("I. INTRODUCTION / EXTRACTED CONTENT", level=2)
    
    for paragraph in text_content.split('\n\n'):
        if paragraph.strip():
            p = doc.add_paragraph(paragraph.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
    # Save
    out_path = "YuvaSaarthi_IEEE_Paper.docx"
    doc.save(out_path)
    print(f"Successfully saved to {out_path}")

if __name__ == "__main__":
    main()
